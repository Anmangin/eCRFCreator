
import json
import xml.etree.ElementTree as ET
import xlsxwriter
from bs4 import BeautifulSoup
from docx.shared import Cm

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from bs4 import BeautifulSoup

i_couleur=0
def add_cover_page(document):
    from datetime import datetime
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor


    # 🎯 Grand espace haut (respiration)
    for _ in range(5):
        document.add_paragraph("")

    # 🧾 Titre principal
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MAQUETTE DE CRF")
    run.bold = True
    run.font.size = Pt(28)

    # 📚 Sous-titre étude
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DE L'ÉTUDE [...]")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # 📏 Espace
    for _ in range(3):
        document.add_paragraph("")

    # 🧬 Ligne séparation soft
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("────────────────────────────")
    run.font.color.rgb = RGBColor(150, 150, 150)

    # 📅 Date génération
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y')}")
    run.italic = True
    run.font.size = Pt(10)

    # 🏢 Signature (optionnel mais pro)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document généré automatiquement")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)


def add_summary_table(graph, document, head):
    document.add_page_break()
    document.add_paragraph("Résumé du CRF", style='Heading 1')

    table = document.add_table(rows=1, cols=2)
    table.width = Cm(20)

    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Niveau"
    hdr_cells[0].width = Cm(1)
    hdr_cells[1].text = "Nom"
    hdr_cells[1].width = Cm(19)

    niveau_map = {
        "ProTrial": ("Trial", 0),
        "ProSite": ("Site", 1),
        "ProPatient": ("Patient", 2),
        "ProVisit": ("Visit", 3),
        "ProForm": ("", 4)  # 👉 vide pour Form
    }

    def recurse(node_id):
        node = graph[node_id]
        tag = node.get('tag', '')
        children = node.get('child', [])

        if tag not in niveau_map:
            for child_id, _ in children:
                recurse(child_id)
            return

        desc = node.get('Description') or node.get('Caption') or ""
        desc = BeautifulSoup(desc, "html.parser").get_text().strip()

        # 👉 repeat
        info = ""
        try:
            if int(node.get('MaxOccurance', 0)) > 1:
                info = "⟳"
        except:
            pass

        # 👉 concat desc + info
        if info:
            desc = f"{desc} - {info}"

        niveau_label, level = niveau_map[tag]

        # 👉 indentation avec tabulation
        indent = "  " * level

        row_cells = table.add_row().cells
        row_cells[0].text = niveau_label
        hdr_cells[0].width = Cm(2)
        row_cells[1].text = indent + desc
        hdr_cells[1].width = Cm(18)

        # 🎨 style léger mais efficace
        if tag == "ProVisit":
            for cell in row_cells:
                cell.paragraphs[0].runs[0].bold = True

        if tag == "ProForm":
            row_cells[1].paragraphs[0].runs[0].italic = True

        for child_id, _ in sorted(children, key=lambda x: x[1]):
            recurse(child_id)

    recurse(head)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')   # type de ligne
        border.set(qn('w:sz'), '6')         # épaisseur (6 = fin)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # noir
        borders.append(border)

    tblPr.append(borders)

def set_cell_background(cell, color="D9EAF7"):  # bleu clair sympa
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    
    tcPr.append(shd)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def remove_details_tags(text):
    # Utilisation de l'expression régulière pour supprimer les balises <details>...</details>
    cleaned_text = re.sub(r'<(th|td)\s+class=["\']check["\'][^>]*>.*?<!--\$htmlbalise-->.', '', text, flags=re.DOTALL)
    # cleaned_text = re.sub(r'<(th|td)\s+class=["\']check["\'][^>]*>.*?<!--\$htmlbalise-->.*?</\1>', '', text, flags=re.DOTALL)


    return cleaned_text

    
def Get_objt(root, racine, keyname, fields):
        ProForm = {}
        for FWAW in root.iter(racine):
            key = FWAW.findtext(keyname)
            if key:  # Vérifie que la clé existe
                ProForm[key] = {field: FWAW.findtext(field) for field in fields}
        return ProForm

def filtrer_par_cle(dictionnaire, champ, valeur):
    return {key: val for key, val in dictionnaire.items() if val.get(champ) == valeur}





def lire_et_trier_donnees(pathfileXML, config_path=r'Python\config.json'):


    # with open("print.txt", "w") as log_file:
    #         log_file.write(config_path)

    # Charger la configuration
    with open(config_path, 'r') as f:

        config = json.load(f)

    # Charger et analyser l'XML
    tree = ET.parse(pathfileXML)
    root = tree.getroot()
    version = root.attrib.get('ver')
    data = {"version": version}

    # Extraire et transformer les données à partir du fichier XML selon la config
    for key, params in config.items():
        racine = params["racine"]
        keyname = params["keyname"]
        fields = params["fields"]
        data[key] = Get_objt(root, racine, keyname, fields)

    # Trier les clés spécifiques


    trier_cles = ["ProPatientVisit", "ProVisitForm", "ProFormGroup", "ProGroupItem", "ProCodeListItem"]
    for key in trier_cles:
        if key in data:
            data[key] = dict(sorted(data[key].items(), key=lambda item: int(item[1].get("OrderNo", 0))))
    # Ajouter le champ "Display" pour ProCodeList
    data["ProCodeList"] = ajouter_display_pro_codelist(data)

    # Ajouter le champ "Display" pour ProItem
    data["ProItem"] = ajouter_display_pro_item(data)

    return data


def ajouter_display_pro_codelist(data):
    """Ajoute un champ 'Display' pour les ProCodeList."""
    for key, value in data["ProCodeList"].items():
        # Filtrer les items associés à la ProCodeList actuelle
        filtered_items = {k: v for k, v in data["ProCodeListItem"].items() if v["ProCodeListGuid"] == key}

        # Construire la chaîne de caractères pour le champ Display
        if len(filtered_items) < 15:
            rep = "<br>".join(f"🔘 {item['Value']} - <b>{item['Caption']}</b>" for item in filtered_items.values())
        else:
            rep = "🔘 Radio bouton trop long"

        # Ajouter le champ Display
        data["ProCodeList"][key]["Display"] = rep

    return data["ProCodeList"]


def ajouter_display_pro_item(data):
    """Ajoute un champ 'Display' pour les ProItem."""
    for key, value in data["ProItem"].items():
            
            if value["ProDataTypeId"] == "5" and value["ProControlTypeId"] == "6":
                print(value)
                rep = "📅 YYYY"
            elif value["ProDataTypeId"] == "5":
                rep = "📅 DD/MM/YYYY"
            else:
                rep = f"{value['SasType']} - {value['MaxLength']}"

            # Ajouter le champ Display
            if rep :data["ProItem"][key]["Display"] = rep





            # Partie du la nature de la var
            if value["Hidden"]=="True": Hidden="👻"
            else: Hidden=""
            if value["ReadOnly"]=="True": ReadOnly="🔒"
            else: ReadOnly=""
            value["Status"]= Hidden + ReadOnly


    return data["ProItem"]

# Appel de la fonction


def find_arbo(TBNode):
            if TBNode["ParentTBNodeId"]=="6":source="Fiche"
            elif TBNode["ParentTBNodeId"]=="7":source="groupe"
            elif TBNode["ParentTBNodeId"]=="8":source="item"
            elif TBNode["ParentTBNodeId"]=="9":source="codelist"
            elif int(TBNode["ParentTBNodeId"]) < 10 : source="unknown"
            else : source="unknown"
            return source
    

def get_check_type(number):
    if   number =="1":return "Valid"
    elif number =="2":return "Required"
    elif number =="3":return "Enabled"
    elif number =="5":return "Required Enabled"
    elif number =="6":return "Hidden"
    elif number =="7":return "Read only"
    elif number =="8":return "Caption change"
    elif number =="9":return "DVC"
    elif number =="10":return "DVA"
    elif number =="11":return "Email"
    elif number =="12":return "Create Form"
    elif number =="14":return "Post save warning"
    elif number =="17":return "Unique value required"
    elif number =="21":return "Disabled Hidden"
    elif number =="22":return "Dynamic codelist setup"
    elif number =="23":return "Dynamic codelist filter"
    return number

def check_data_type(number):
    if   number =="1":return "Text Box"
    elif number =="2":return "Text Area"
    elif number =="3":return "Radio Button"
    elif number =="4":return "Check Box"
    elif number =="5":return "Combo Box"
    elif number =="6":return "Year"
    elif number =="7":return "YearMonth"
    elif number =="8":return "Date"
    elif number =="10":return "Time"
    elif number =="13":return "Label"
    elif number =="16":return "Textbox with MedDRA"
    elif number =="18":return "Timer"
    elif number =="21":return "Combo Box Dynamic"
    return number


def get_message(data,ProItemGuid,ProGroupGuid,ProFormGuid):
    """genere le message pour les editcheck"""
    Message=""
    i=0
    ProEdit=filtrer_par_cle( data,"ProItemGuid",ProItemGuid)
    for  key,Edit in ProEdit.items():
        if (Edit["ProItemGuid"] == ProItemGuid) and (Edit["TargetProFormGuid"]==ProFormGuid or Edit["TargetProFormGuid"] is None) and (Edit["TargetProGroupGuid"]==ProGroupGuid or Edit["TargetProGroupGuid"] is None)  :
            i+=1
            EditACtion= Edit["ProEditActionId"]
            EditACtion = get_check_type(EditACtion)
            msg=Edit["Message"]
            chk=Edit["ActionExpression"]
            DTE=Edit['DataExpression']
            path=Edit["TargetPath"]

            Message+=f"<tr><td> {EditACtion}:{path}</td> </tr>"
            Message+=f"<tr> <td> <pre><code class='javascript'>#Action Expression \n{chk} \n#data Expression \n{DTE} \n</code></pre> </td>"
            Message+=f"<td> {msg}</td> </tr>"


    Message+="</table></details>"
    if i==0:Message=""
    else: Message= f"<details> <summary>{i} EditCheck </summary><table>" + Message


    return Message




def exporter_donnees_markdown_eCRF(data,unic_form):

        order=0
        JSON_EXPORT = {}
        for key, value in data["ProPatient"].items():
            print(value)
            ProFormGuid=value["ProFormGuid"]
            Caption=value["Caption"]
            order,JSON_EXPORT = ADD_FORM(data,"Patient Template",ProFormGuid,order,JSON_EXPORT)
       


        for  PVkey,PatientVisit in data["ProPatientVisit"].items():



            # Ecriture des variables nécéssaire pour l'affichage
            ProVisitGuid        =PatientVisit["ProVisitGuid"]
            V_description       =data["ProVisit"][ProVisitGuid]["Description"]
            V_Caption           =data["ProVisit"][ProVisitGuid]["Caption"]
            V_OrderNo           =PatientVisit["OrderNo"]



            ProVisitForm=filtrer_par_cle( data["ProVisitForm"],"ProVisitGuid",ProVisitGuid)
           
            for  VFkey,VisitForm in ProVisitForm.items():
                ProFormGuid=VisitForm.get("ProFormGuid")
                written=data["ProForm"][ProFormGuid].get("written", False)
                if written and unic_form: continue    

                # Ecriture des variables nécéssaire pour l'affichage
                F_OrderNo           =VisitForm["OrderNo"]

                data["ProForm"][ProFormGuid]["written"]=True
                order,JSON_EXPORT = ADD_FORM(data,V_description,ProFormGuid,order,JSON_EXPORT)

        return JSON_EXPORT
        

def ADD_FORM(data,V_description,ProFormGuid,order,JSON_EXPORT):

                F_description       =data["ProForm"][ProFormGuid]["Description"]
                F_SasName           =data["ProForm"][ProFormGuid]["SasName"]
                F_Caption           =data["ProForm"][ProFormGuid]["Caption"]
    
                ProFormGroup=filtrer_par_cle( data["ProFormGroup"],"ProFormGuid",ProFormGuid)
                for  key,FormGroup in ProFormGroup.items():
                    # Ecriture des variables nécéssaire pour l'affichage
                    ProGroupGuid        =FormGroup["ProGroupGuid"]
                    G_description       =data["ProGroup"][ProGroupGuid]["Caption"]
                    G_Caption           =data["ProGroup"][ProGroupGuid]["Caption"]

                    ProGroupItem=filtrer_par_cle( data["ProGroupItem"],"ProGroupGuid",ProGroupGuid)

                    for  GI_key,GroupItem in ProGroupItem.items():
                        # Ecriture des variables nécéssaire pour l'affichage
                        I_OrderNo           =GroupItem["OrderNo"]
                        ProItemGuid         =GroupItem.get("ProItemGuid")
                        I_description       =data["ProItem"][ProItemGuid]["Description"]
                        i_SasName           =data["ProItem"][ProItemGuid]["SasName"]
                        I_Caption           =data["ProItem"][ProItemGuid]["Caption"]
                        ProCodeListGuid     =data["ProItem"][ProItemGuid]["ProCodeListGuid"]
                        Hidden              =data["ProItem"][ProItemGuid]["Hidden"]
                        Disabled            =data["ProItem"][ProItemGuid]["Disabled"]
                        ReadOnly            =data["ProItem"][ProItemGuid]["ReadOnly"]
                        i_Display           =data["ProItem"][ProItemGuid]["Display"]
                        I_Status            =data["ProItem"][ProItemGuid]["Status"]
                        ProItemCategoryGuid            =data["ProItem"][ProItemGuid]["ProItemCategoryGuid"]
                        if ProItemCategoryGuid: Cat_Description            =data["ProItemCategory"][ProItemCategoryGuid]["Description"]
                        else :Cat_Description=""
                        rep=""
                        if ProCodeListGuid :rep=data["ProCodeList"][ProCodeListGuid]["Display"]
                      
                        
                        Message = get_message(data["ProEdit"],ProItemGuid,ProGroupGuid,ProFormGuid)    
                        order+=1
                        if Cat_Description != "":Message= "Catérogie:" + Cat_Description + "<br>" +  Message
                        JSON_EXPORT = get_JSONLIGNE(JSON_EXPORT,V_description,F_description,G_description,GI_key,order,I_description,I_Caption,Message,rep,i_Display,I_Status,i_SasName)            
                return  order,JSON_EXPORT




def get_JSONLIGNE(JSON_EXPORT, V_description, F_description, G_description, GI_key,
                  order,I_description,I_caption, Message, rep,display, I_Status, i_SasName):
    # Assure-toi que la clé 'Patient' existe dans JSON_EXPORT
    if 'visites' not in JSON_EXPORT:
        JSON_EXPORT['visites'] = []

    # Cherche si le patient existe déjà dans la liste des patients, sinon crée-le
    visites = next((p for p in JSON_EXPORT['visites'] if p['V_description'] == V_description), None)
    
    if not visites:
        # Crée un nouveau patient avec V_description comme nom
        visites = {"V_description": V_description, "fiches": []}
        JSON_EXPORT['visites'].append(visites)

    # Cherche la fiche dans le patient, sinon crée-la
    fiche = next((f for f in visites['fiches'] if f['F_description'] == F_description), None)
    
    if not fiche:
        fiche = {"F_description": F_description, "groupes": []}
        visites['fiches'].append(fiche)

    # Cherche le groupe dans la fiche, sinon crée-le
    groupe = next((g for g in fiche['groupes'] if g['G_description'] == G_description), None)
    
    if not groupe:
        groupe = {"G_description": G_description, "questions": []}
        fiche['groupes'].append(groupe)


    # Ajout de la question dans le groupe
    question = {
         "GI_key": GI_key,
         "Item_Order":order,
        "I_description": I_description,
        "I_caption":I_caption,
        "Message": Message if Message else None,
        "rep": rep if rep else display,
        "I_Status": I_Status,
        "i_SasName": i_SasName
    }
    groupe['questions'].append(question)
    
    return JSON_EXPORT

def find_root(graph):
    for node_id, node in graph.items():
        if not node['parent']:  # pas de parent
            return node_id

def create_graph(data):
    dictionnary = dict()
    nodes = ["Trial", "Site", "Patient", "Visit", "Form", "Group", "Item", "CodeList", "CodeListItem"]
    for nodetype in nodes:
        name = "Pro" + nodetype
        for elem in data.iter(name):
            try:
                guid = elem.find("ProObjectGuid").text
                dictionnary[guid] = dict()
                dictionnary[guid]['child'] = list()
                dictionnary[guid]['parent'] = list()
                caption = elem.find('Caption').text if elem.find('Caption') is not None else ""
                value_elem = elem.find('Value')

                if value_elem is not None and value_elem.text:
                    dictionnary[guid]['Caption'] = f"{value_elem.text} - {caption}"
                else:
                    dictionnary[guid]['Caption'] = caption
                dictionnary[guid]['tag'] = elem.tag
            except AttributeError:
                pass
            try:
                dictionnary[guid]['SasName'] = elem.find('SasName').text
            except:
                pass
            try:
                dictionnary[guid]['type'] = elem.find('ProControlTypeId').text
            except:
                pass
            try:
                if name != "ProCodeListItem":
                    dictionnary[guid]['child'].append((elem.find('ProCodeListGuid').text, 1))
            except:
                pass
            try:
                dictionnary[elem.find('ProCodeListGuid').text]['parent'].append(guid)
            except:
                pass
            try:
                dictionnary[guid]['Description'] = elem.find('Description').text
            except:
                pass
    # Links
    for i in range(6):
        link = "Pro" + nodes[i] + nodes[i+1]
        for elem in data.iter(link):
            parent = elem.find("Pro" + nodes[i]  + "Guid").text
            child =  elem.find("Pro" + nodes[i+1]+ "Guid").text
            order = int(elem.find("OrderNo").text)
            dictionnary[parent]['child'].append((child, order))
            dictionnary[child]['parent'].append(parent)
            try:
                dictionnary[child]['MaxOccurance'] = elem.find('MaxOccurance').text
            except:
                pass
    for elem in data.iter("ProCodeListItem"):
        parent = elem.find("ProCodeListGuid").text
        child = elem.find("ProObjectGuid").text
        order = int(elem.find("OrderNo").text)
        dictionnary[parent]['child'].append((child, order))
        dictionnary[child]['parent'].append(parent)
        dictionnary[guid]['Caption'] = elem.find('Caption').text
        dictionnary[guid]['Value'] = elem.find('Value').text
    return dictionnary

def create_edit_check_dictionnary(data):
    dictionnary = dict()
    keys = ["OID", "ProEditActionId", "TargetPath", "ActionExpression", "DataExpression"]
    for proedit in data.iter("ProEdit"):
        guids = list()
        try:
            guids.append(proedit.find("ProItemGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProGroupGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProFormGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProVisitGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProPatientGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProSiteGuid").text)
        except AttributeError:
             pass
        try:
            guids.append(proedit.find("ProTrialGuid").text)
        except:
             pass
        OID = proedit.find('OID').text
        dictionnary[OID] = dict()
        if len(guids) == 0: print(OID)
        for k in keys:
            sub = proedit.find(k).text
            dictionnary[OID][k] = sub
        dictionnary[OID]["lower"] = guids[0]
    return dictionnary

# Useless, does not function with return
def recursive_move(graph, elem):
    try:
        print (graph[elem]['Caption'], graph[elem]['SasName'], check_data_type(graph[elem]['type']))
    except:
        loc = graph[elem]['child']
        for e in loc:
            recursive_move(graph, e)

def print_xls_from_edit_check(xmlFile, xlsName):
    data = ET.parse(xmlFile)
    editChecks = create_edit_check_dictionnary(data)
    graph = create_graph(data)
    # Create an new Excel file and add a worksheet
    workbook = xlsxwriter.Workbook(xlsName)
    worksheet = workbook.add_worksheet("Checks")
    cell_format = workbook.add_format()
    cell_format.set_text_wrap()
    ROWS = ['Item Name', "Item Type", "Test Type", "Action", "Data"]
    # TITLE ROW
    for i in range(len(ROWS)):
         worksheet.write(0, i, ROWS[i])
    l = 1
    checked = list()
    for k, v in editChecks.items():
        action = v['ActionExpression']
        # Remove the empty lines
        try:
            action = "\n".join([ll.rstrip() for ll in action.splitlines() if ll.strip()])
        except AttributeError:
             action = ""
        data = v['DataExpression']
        # Remove the empty lines
        try:
            data = "\n".join([ll.rstrip() for ll in data.splitlines() if ll.strip()])
        except AttributeError:
            data = ""           
        item = graph[editChecks[k]['lower']]
        checked.append(editChecks[k]['lower'])
        worksheet.write(l, 0, item['Caption'])
        worksheet.write(l, 1,item['tag'])
        worksheet.write(l, 2, get_check_type(v['ProEditActionId']))
        worksheet.write(l, 3, action, cell_format)
        worksheet.write(l, 4, data, cell_format)
        l += 1
    worksheet.autofit()
    worksheet.autofilter(0, 0, l - 1, len(ROWS)-1)
    worksheet2 = workbook.add_worksheet('Items Without Checks')
    l = 1
    worksheet2.write(0, 0, 'Caption', cell_format)
    worksheet2.write(0, 1, 'SasName', cell_format)
    worksheet2.write(0, 2, 'type', cell_format)
    for id, node in graph.items():
        if node['tag'] == "ProItem":
            if id not in checked:
                worksheet2.write(l, 0, node['Caption'], cell_format)
                worksheet2.write(l, 1, node['SasName'], cell_format)
                worksheet2.write(l, 2, check_data_type(node['type']), cell_format)
                l += 1
    worksheet2.autofilter(0, 0, l - 1, 2)
    worksheet2.autofit()
    workbook.close()


def print_graph(graph, head, lvl=0, fileName=None):
    lns = graph[head]['child']
    tag = graph[head]['tag']
    try:
        desc = graph[head]['Description'].strip()
        try:
            typ = check_data_type(graph[head]['type'])
        except:
            typ = ''
        try:
            sasNam = graph[head]['SasName']
        except:
            sasNam = ''
        if len(desc) > 0:
            print(desc, typ, sep="\t", file=fileName)
    except:
        print("\t", graph[head]['Caption'], file=fileName)
    for e in sorted(lns, key = lambda e:e[1]):
        lvl = lvl + 1
        print_graph(graph, e[0], lvl, fileName)


def print_doc_xml(xmlFile, docFile, head=None):
    data = ET.parse(xmlFile)
    graph = create_graph(data)
    # Find the head
    if head == None:
        for k, v in graph.items():
            if v['tag'] == 'ProPatient':
                head = k
                break
    from docx import Document
    document = Document()
    # document.add_heading("SUMMARY", level=1)
    # internal_func_doc(graph, document, head=head, lvl=2, unique=False, summary=True)
    # document.add_page_break()
    # document.add_heading("CRF", level=1)

    # # Créer un style basé sur Titre1
    styles = document.styles
    if 'FormDescription' not in styles:
        style = styles.add_style('FormDescription', 1)  # 1 = paragraph style
        style.base_style = styles['Heading 1']          # Hérite de Heading 1
        font = style.font
        font.name = 'Arial'                             # Nom de police
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0, 51, 102)           # Bleu foncé

        # head = find_root(graph) # More than one node without parent

        add_cover_page(document)
        add_summary_table(graph, document, head)
        document.add_page_break()
        # Ajoute la fiche PAT au début du document 
        PAT = None
        for k, v in graph.items():
            if v['tag'] == 'ProForm' and v['Caption'] == "PAT":
                PAT = k
                break
        internal_func_doc(graph, document, head=PAT, lvl=2, buffer=list())
        # Ajoute les fiches descendants de head
        internal_func_doc(graph, document, head=head, lvl=2, buffer=list())
        document.save(docFile)

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color="D9EAF7"):
    """Applique un fond coloré à une cellule Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)





def internal_func_doc(graph, documentName, head=None, lvl=0, buffer=list(), maxi=False, unique=True,uniqueList=list(), summary=False):
    lns = graph[head]['child']
    tag = graph[head]['tag']
    if tag == 'ProVisit' and not summary:
        documentName.add_page_break()
    if summary and tag not in ["ProTrial", "ProSite", "ProPatient", "ProVisit", "ProForm"]:
        return
    try:
        desc = graph[head]['Description'].strip()
        soup = BeautifulSoup(desc, "html.parser")
        desc = soup.get_text()
        if unique and tag in ["ProTrial", "ProSite", "ProPatient", "ProVisit", "ProForm"]:
            if head in uniqueList: return
            uniqueList.append(head)
        try:
            typ = check_data_type(graph[head]['type'])
        except:
            typ = ''
        try:
            hidden = graph[head]['Hidden']
        except:
            hidden = False
        try:
            sasNam = graph[head]['SasName']
        except:
            sasNam = ''
        try:
            repeat = graph[head]['MaxOccurance']
            if int(repeat) > 1:
                desc += " ⟳"
        except:
            repeat = ""
        if len(desc) > 0:
            if len(typ) == 0:
                try:
                    documentName.add_heading(desc, level=lvl)
                except ValueError:
                    documentName.add_heading(desc, level=9)
            else:
                buffer = (desc, typ, list(), hidden)
    except KeyError:
        buffer[2].append(graph[head]['Caption'])
    if len(buffer) > 0 and len(lns) == 0 and (maxi or tag != 'ProCodeListItem'):
        table = documentName.add_table(rows=1, cols=2)
        if buffer[3] and buffer[3] == "True":
            table.style = 'Medium Shading 1'
        else:
            table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].width = Cm(14)
        hdr_cells[0].text = buffer[0].replace("\\r\\n", "\n")
        codeList = "\n".join(buffer[2])
        hdr_cells[1].width = Cm(6)
        if len(buffer[2]) > 10:
            codeList = "\n".join(buffer[2][0:4] + ['...'] + buffer[2][-4:-1])
        if len(codeList)>0:
            para = hdr_cells[1].paragraphs[0].add_run(codeList)
            para.font.color.rgb = RGBColor(255, 0, 0)
        else:
            hdr_cells[1].text = buffer[1]
        buffer = list()
    ls = sorted(lns, key = lambda e:e[1])
    lvl = lvl + 1
    for i in range(len(ls)):
        e = ls[i]
        if i == len(ls) - 1:
            maxi = True
        else:
            maxi = False
        internal_func_doc(graph, documentName, e[0], lvl, buffer, maxi, unique, uniqueList, summary=summary)



def find_parentals(graph, id):
    for e in graph[id]['parents']:
        print(e)
        find_parentals(graph, e)

# def internal_func_doc(graph, documentName, head=None, lvl=0, buffer=list(), maxi=False):
#     lns = graph[head]['child']
#     tag = graph[head]['tag']
    
#     try:
#         desc = graph[head]['Description'].strip()
#         soup = BeautifulSoup(desc, "html.parser")
#         desc = soup.get_text()
        
#         try:
#             typ = check_data_type(graph[head]['type'])
#         except:
#             typ = ''
#         try:
#             sasNam = graph[head]['SasName']
#         except:
#             sasNam = ''
#         try:
#             repeat = graph[head]['MaxOccurance']
#             if int(repeat) > 1:
#                 desc += " 🔁:" + repeat
#         except:
#             repeat = ""
        
#         if len(desc) > 0:
#             # Si c'est une fiche, on centre le titre
#             if tag == 'ProForm':
#                 documentName.add_page_break()
#                 para = documentName.add_paragraph(desc, style='FormDescription')

#                 para.alignment = WD_ALIGN_PARAGRAPH.CENTER
#                 run = para.runs[0]
#                 run.bold = True
#                 run.font.size = Pt(16)
#             # Si c'est un groupe, on stylise le titre
#             elif tag == 'ProGroup':
#                 para = documentName.add_paragraph(desc)
#                 para.alignment = WD_ALIGN_PARAGRAPH.LEFT
#                 run = para.runs[0]
#                 run.bold = True
#                 run.font.size = Pt(12)
#                 run.font.color.rgb = RGBColor(0, 51, 102)  # bleu foncé
#             else:
#                 buffer = (desc, typ, list())
#     except:
#         buffer[2].append(graph[head]['Caption'])
    
#     # Si on est au bout d'une branche et qu'on doit écrire le tableau
#     if len(buffer) > 0 and len(lns) == 0 and (maxi or tag != 'ProCodeListItem'):
#         table = documentName.add_table(rows=1, cols=2)
#         set_table_borders(table)
#         table.width = Cm(20)
#         row = table.rows[0]
#         prevent_row_split(row)
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = buffer[0].replace("\\r\\n", "\n")
#         hdr_cells[0].width = Cm(13)
#         codeList = "\n".join(buffer[2]) if buffer[2] else ""
#         if codeList and len(buffer[2]) > 10:
#             codeList = "\n".join(buffer[2][0:4] + ['...'] + buffer[2][-4:-1])
#         # Version correcte en Python
#         if codeList:
#             display_value = codeList
#         elif buffer[1] == 'Date':
#             display_value = "📅 DD/MM/YYYY"
#         else:
#             display_value = buffer[1]
#         hdr_cells[1].text = display_value
#         hdr_cells[1].width = Cm(5)
#         hdr_cells[1].alignment = WD_ALIGN_PARAGRAPH.RIGHT

#         cell = hdr_cells[1]
#         cell.text = display_value
#         cell.width = Cm(5)

#         # ⚡ alignement sur le paragraphe interne
#         cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
#         for i, row in enumerate(table.rows):
#             global i_couleur 
#             i_couleur+=1
#             color = "D9EAF7" if i_couleur % 2 == 0 else "FFFFFF"  # bleu clair / blanc
#             for cell in row.cells:
#                 set_cell_background(cell, color=color)

#         buffer = list()
    
#     # Parcours récursif
#     ls = sorted(lns, key=lambda e: e[1])
#     for i, e in enumerate(ls):
#         internal_func_doc(graph, documentName, e[0], lvl + 1, buffer, maxi=(i == len(ls) - 1))
    
